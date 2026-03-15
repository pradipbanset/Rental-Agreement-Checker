# """
# Document Service - Database-backed version
# """
# import PyPDF2
# import io
# from docx import Document as DocxDocument
# import uuid
# from typing import Optional
# from fastapi import UploadFile, HTTPException
# from sqlalchemy.orm import Session

# from app.models.db_models import Document, DocumentStatus

# class DocumentService:
    
#     @staticmethod
#     def detect_state_from_text(text: str) -> str:
#         """Auto-detect state/territory from contract text"""
#         text_lower = text.lower()
        
#         if "act standard tenancy" in text_lower or "residential tenancies act 1997" in text_lower:
#             return "ACT"
#         elif "residential tenancies act 2010" in text_lower:
#             return "NSW"
#         elif "nsw" in text_lower and "residential" in text_lower:
#             return "NSW"
#         elif "victoria" in text_lower:
#             return "VIC"
#         elif "queensland" in text_lower:
#             return "QLD"
#         elif "south australia" in text_lower:
#             return "SA"
#         elif "western australia" in text_lower:
#             return "WA"
#         elif "tasmania" in text_lower:
#             return "TAS"
#         elif "northern territory" in text_lower:
#             return "NT"
        
#         return "NSW"  # Default
    
#     @staticmethod
#     async def extract_text_from_file(file: UploadFile) -> str:
#         """Extract text from uploaded file"""
#         content = await file.read()
#         text = ""
        
#         if file.filename.endswith('.pdf'):
#             try:
#                 pdf_file = io.BytesIO(content)
#                 pdf_reader = PyPDF2.PdfReader(pdf_file)
                
#                 if pdf_reader.is_encrypted:
#                     raise HTTPException(status_code=400, detail="PDF is password protected")
                
#                 for page in pdf_reader.pages:
#                     page_text = page.extract_text()
#                     if page_text:
#                         text += page_text + "\n"
                
#                 if not text.strip():
#                     raise HTTPException(status_code=400, detail="Could not extract text from PDF.")
                    
#             except PyPDF2.errors.PdfReadError as e:
#                 raise HTTPException(status_code=400, detail=f"Invalid PDF: {str(e)}")
        
#         elif file.filename.endswith('.docx'):
#             doc_file = io.BytesIO(content)
#             doc = DocxDocument(doc_file)
#             text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
#         elif file.filename.endswith('.txt'):
#             try:
#                 text = content.decode('utf-8')
#             except UnicodeDecodeError:
#                 text = content.decode('latin-1')
        
#         else:
#             raise HTTPException(status_code=400, detail="Unsupported file type")
        
#         if not text.strip():
#             raise HTTPException(status_code=400, detail="File is empty")
        
#         return text
    
#     @staticmethod
#     def create_document(db: Session, filename: str, text: str, state: str) -> Document:
#         """Create and store document in database"""
#         doc_id = str(uuid.uuid4())
#         detected_state = DocumentService.detect_state_from_text(text)
#         final_state = state if state != "NSW" else detected_state
        
#         print(f"✓ Uploaded {filename}")
#         print(f"  Detected: {detected_state} | Selected: {state} | Final: {final_state}")
        
#         # Create document
#         document = Document(
#             id=doc_id,
#             filename=filename,
#             state=final_state,
#             detected_state=detected_state,
#             text=text,
#             status=DocumentStatus.UPLOADED
#         )
        
#         db.add(document)
#         db.commit()
#         db.refresh(document)
        
#         return document
    
#     @staticmethod
#     def get_document(db: Session, doc_id: str) -> Document:
#         """Get document by ID"""
#         document = db.query(Document).filter(Document.id == doc_id).first()
#         if not document:
#             raise HTTPException(status_code=404, detail="Document not found")
#         return document
    
#     @staticmethod
#     def get_latest_completed_document(db: Session) -> Optional[Document]:
#         """Get the latest completed document"""
#         return db.query(Document).filter(
#             Document.status == DocumentStatus.COMPLETED
#         ).order_by(Document.updated_at.desc()).first()
    
#     @staticmethod
#     def update_status(db: Session, doc_id: str, status: DocumentStatus):
#         """Update document status"""
#         document = DocumentService.get_document(db, doc_id)
#         document.status = status
#         db.commit()
    
#     @staticmethod
#     def set_error(db: Session, doc_id: str, error: str):
#         """Mark document as errored"""
#         document = DocumentService.get_document(db, doc_id)
#         document.status = DocumentStatus.ERROR
#         document.error_message = error
#         db.commit()

# # Singleton instance
# document_service = DocumentService()






"""
Document Service - Database-backed version
"""
import PyPDF2
import io
from docx import Document as DocxDocument
import uuid
from typing import Optional
from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session

from app.models.db_models import Document, DocumentStatus

class DocumentService:
    
    @staticmethod
    def detect_state_from_text(text: str) -> str:
        """Auto-detect state/territory from contract text"""
        text_lower = text.lower()
        
        if "act standard tenancy" in text_lower or "residential tenancies act 1997" in text_lower:
            return "ACT"
        elif "residential tenancies act 2010" in text_lower:
            return "NSW"
        elif "nsw" in text_lower and "residential" in text_lower:
            return "NSW"
        elif "victoria" in text_lower:
            return "VIC"
        elif "queensland" in text_lower:
            return "QLD"
        elif "south australia" in text_lower:
            return "SA"
        elif "western australia" in text_lower:
            return "WA"
        elif "tasmania" in text_lower:
            return "TAS"
        elif "northern territory" in text_lower:
            return "NT"
        
        return "NSW"  # Default
    
    @staticmethod
    async def extract_text_from_file(file: UploadFile) -> str:
        """Extract text from uploaded file"""
        content = await file.read()
        text = ""
        
        if file.filename.endswith('.pdf'):
            try:
                pdf_file = io.BytesIO(content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                if pdf_reader.is_encrypted:
                    raise HTTPException(status_code=400, detail="PDF is password protected")
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                if not text.strip():
                    raise HTTPException(status_code=400, detail="Could not extract text from PDF.")
                    
            except PyPDF2.errors.PdfReadError as e:
                raise HTTPException(status_code=400, detail=f"Invalid PDF: {str(e)}")
        
        elif file.filename.endswith('.docx'):
            doc_file = io.BytesIO(content)
            doc = DocxDocument(doc_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        elif file.filename.endswith('.txt'):
            try:
                text = content.decode('utf-8')
            except UnicodeDecodeError:
                text = content.decode('latin-1')
        
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="File is empty")
        
        return text
    
    @staticmethod
    def create_document(db: Session, filename: str, text: str, state: str, user_id: int = None) -> Document:
        """Create and store document in database"""
        doc_id = str(uuid.uuid4())
        detected_state = DocumentService.detect_state_from_text(text)
        final_state = state if state != "NSW" else detected_state
        
        print(f"\n📝 Creating document in DB:")
        print(f"   Filename: {filename}")
        print(f"   Text length: {len(text)}")
        print(f"   Detected: {detected_state} | Selected: {state} | Final: {final_state}")
        print(f"   User ID: {user_id}")
        
        # Create document
        document = Document(
            id=doc_id,
            filename=filename,
            state=final_state,
            detected_state=detected_state,
            text=text,
            status=DocumentStatus.UPLOADED,
            user_id=user_id
        )
        
        print(f"   Document object created, adding to session...")
        db.add(document)
        
        print(f"   Committing to database...")
        db.commit()
        
        print(f"   Refreshing object...")
        db.refresh(document)
        
        print(f"   ✓ Document saved successfully")
        print(f"   Text length after refresh: {len(document.text) if document.text else 0}")
        
        return document
    
    @staticmethod
    def get_document(db: Session, doc_id: str) -> Document:
        """Get document by ID"""
        document = db.query(Document).filter(Document.id == doc_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        return document
    
    @staticmethod
    def get_latest_completed_document(db: Session) -> Optional[Document]:
        """Get the latest completed document"""
        return db.query(Document).filter(
            Document.status == DocumentStatus.COMPLETED
        ).order_by(Document.updated_at.desc()).first()
    
    @staticmethod
    def update_status(db: Session, doc_id: str, status: DocumentStatus):
        """Update document status"""
        document = DocumentService.get_document(db, doc_id)
        document.status = status
        db.commit()
    
    @staticmethod
    def set_error(db: Session, doc_id: str, error: str):
        """Mark document as errored"""
        document = DocumentService.get_document(db, doc_id)
        document.status = DocumentStatus.ERROR
        document.error_message = error
        db.commit()

# Singleton instance
document_service = DocumentService()